import streamlit as st  # Pentru interfața web
import spacy  # Pentru procesarea limbajului natural (tokenizare, embedding-uri)
import glob  # Pentru navigarea în fișiere
import os  # Pentru operații cu căi de fișiere
import numpy as np  # Pentru operații matematice
from docx import Document  # Pentru citirea fișierelor Word (.docx)
from sklearn.feature_extraction.text import TfidfVectorizer  # Vectorizare text
from sklearn.metrics.pairwise import cosine_similarity  # Calcul similaritate
from sklearn.preprocessing import MinMaxScaler  # Normalizare scoruri
from sentence_transformers import SentenceTransformer, util  # Embedding-uri avansate
import json  # Pentru citirea/scrierea scorurilor precalculate
import re  # Pentru căutare text cu regex
import warnings
warnings.filterwarnings('ignore', category=UserWarning)

nlp = spacy.load("en_core_web_lg")


def load_docx_from_folder(folder_path, is_cv=True):
    """
    Încarcă toate fișierele .docx dintr-un folder și extrage textul.
    
    Args:
        folder_path (str): Calea către folder (ex: './DataSet/cv')
        is_cv (bool): Dacă True, se procesează ca CV; False pentru job descriptions
    
    Returns:
        tuple: (lista cu texte, lista cu nume de fișiere, lista cu descrieri)
    """
    documents = []
    filenames = []
    descriptions = []
    
    for filepath in glob.glob(os.path.join(folder_path, '*.docx')):
        # Extrage textul din fiecare paragraf
        text = extract_text_from_docx(filepath, is_cv)
        
        # Curăță spații și linii noi
        text = text.replace('\n', ' ').replace('  ', ' ')
        
        # Procesare specifică
        if not is_cv:  # Pentru job descriptions
            text = text.split("Benefits:")[0]  # Ignoră secțiunea Benefits
            description = text.split("Key Responsibilities:")[0]  # Extrage descrierea principală
        else:  # Pentru CV-uri
            description = text.split("Project Experience")[1] if "Project Experience" in text else ""
        
        documents.append(text)
        filenames.append(os.path.basename(filepath))
        descriptions.append(description)
    
    return documents, filenames, descriptions

def extract_text_from_docx(filepath, is_cv):
    """
    Extrage textul dintr-un fișier .docx, ignorând primul paragraf pentru CV-uri.
    
    Args:
        filepath (str): Calea către fișier
        is_cv (bool): Dacă True, ignoră primul paragraf (presupus a fi header)
    
    Returns:
        str: Textul concatenat din toate paragrafele
    """
    doc = Document(filepath)
    paragraphs = doc.paragraphs[1:] if is_cv else doc.paragraphs  # Ignoră header-ul pentru CV
    return ' '.join(para.text for para in paragraphs)

# Custom tokenizer using spaCy
def spacy_tokenizer(text):
    """
    Tokenizează textul folosind spaCy, cu filtrare avansată.
    
    Args:
        text (str): Textul de procesat
    
    Returns:
        list: Liste de tokeni lematizați și filtrați
    """
    doc = nlp(text)
    return [
        token.lemma_.lower()  # Reduce cuvintele la forma de bază
        for token in doc
        if not token.is_stop and not token.is_punct  # Ignora stopwords și punctuatie
        and not token.text.isspace() and not token.text.isnumeric()  # Ignora spatii/numere
    ]

def progress_bar_update(percent_complete, 
                        progress_bar,
                        status_text):
    progress_bar.progress(percent_complete)
    progress_text = "Operation in progress. Please wait. ⏳" if percent_complete < 100 else "Operation Completed."
    status_text.text(f"{progress_text} - {percent_complete}% - ")

