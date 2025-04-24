import streamlit as st
import spacy
import glob
import os
import numpy as np
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import MinMaxScaler
from sentence_transformers import SentenceTransformer, util
import json
import re
from utils.file_loader import load_docx_from_folder

# Set your folders here
cv_folder = './DataSet/cv'
job_folder = './DataSet/job_descriptions'

# 1------------------------------
domain_data = {
    "Banking": { "desc": "Developed software for banking platforms, digital wallets, loan processing, or credit systems.",
                 "keywords": ["bank", "loan", "credit", "atm", "fintech", "interest", "account", "ledger"] },
    "Healthcare": { "desc": "Built medical systems such as EHR, hospital platforms, clinical apps, or telemedicine services.",
                    "keywords": ["healthcare", "ehr", "patient", "hospital", "clinic", "medical", "doctor", "nurse"] },
    "E-commerce": { "desc": "Created platforms or tools for online stores, shopping carts, payments, or product discovery.",
                    "keywords": ["ecommerce", "checkout", "cart", "payment", "shopify", "woocommerce", "product", "sku"] },
    "Telecommunications": { "desc": "Engineered tools for telecom networks, call management, VoIP, or network monitoring.",
                            "keywords": ["telecom", "sms", "voip", "5g", "network", "bandwidth", "subscriber", "lte"] },
    "Education": { "desc": "Built education platforms such as learning portals, student dashboards, or LMS systems.",
                   "keywords": ["education", "student", "teacher", "learning", "course", "classroom", "lms", "school"] },
    "Retail": { "desc": "Developed software for retail businesses such as POS systems, inventory, or loyalty programs.",
                "keywords": ["retail", "store", "inventory", "pos", "stock", "sku", "receipt", "shopping"] },
    "Insurance": { "desc": "Created applications for policy management, claims processing, or underwriting systems.",
                   "keywords": ["insurance", "claim", "policy", "underwriting", "premium", "broker", "risk"] },
    "Legal": { "desc": "Built tools for legal case management, document processing, or e-discovery platforms.",
               "keywords": ["legal", "law", "contract", "case", "compliance", "jurisdiction", "litigation"] },
    "Manufacturing": { "desc": "Developed automation systems, supply chain tools, or MES solutions for production plants.",
                       "keywords": ["manufacturing", "plant", "automation", "mes", "machine", "factory", "assembly"] },
    "Transportation & Logistics": { "desc": "Built platforms for delivery tracking, fleet management, logistics optimization, or routing.",
                                    "keywords": ["logistics", "delivery", "routing", "fleet", "dispatch", "transport", "warehouse"] },
    "Energy & Utilities": { "desc": "Developed monitoring systems, SCADA platforms, or analytics for power and water utilities.",
                            "keywords": ["energy", "power", "electricity", "gas", "grid", "meter", "solar", "utility", "scada"] },
    "Real Estate": { "desc": "Engineered platforms for property listings, CRM tools for agents, or real estate analytics.",
                     "keywords": ["real estate", "property", "mortgage", "agent", "listing", "tenant", "lease"] },
    "Government": { "desc": "Built public service portals, civic data dashboards, or digital identity platforms.",
                    "keywords": ["government", "municipal", "civic", "permit", "id", "citizen", "registry"] },
    "Marketing": { "desc": "Built marketing analytics tools, email campaign platforms, or digital ad performance systems.",
                   "keywords": ["marketing", "campaign", "seo", "email", "ads", "promotion", "branding", "targeting"] },
    "Media & Entertainment": { "desc": "Created digital content platforms, streaming services, or entertainment production tools.",
                               "keywords": ["media", "streaming", "video", "music", "entertainment", "broadcast", "subscriber"] },
    "Construction": { "desc": "Engineered project management tools, BIM integrations, or field apps for construction teams.",
                      "keywords": ["construction", "site", "blueprint", "project", "bim", "architect", "contractor"] },
    "Finance (non-banking)": { "desc": "Worked on accounting systems, budgeting tools, payroll, or financial planning platforms.",
                               "keywords": ["finance", "accounting", "budget", "payroll", "invoice", "expense", "audit", "report"] }
}

#SBERT model for semantic similarity + a dictionary with industries (e.g. banking, finance) + description + keywords
def match_domains(model, domain_data, text, top_k=5, keyword_boost=0.1):
    domain_names = list(domain_data.keys())
    domain_embeddings = model.encode([domain_data[d]["desc"] for d in domain_names], convert_to_tensor=True)

    text_embedding = model.encode(text, convert_to_tensor=True)

    #the embedding is obtained for each industry and the cosine similarity between them is calculated
    cos_scores = util.cos_sim(text_embedding, domain_embeddings)[0]

    #Keyword matching boost
    scores = []
    text_lower = text.lower()
    for i, domain in enumerate(domain_names): #manually search for how many keywords specific to each industry appear in the CV and to the semantic score,
        # we add a penalty/reward for each keyword found.
        # => Compensates when embedding doesn't catch small or specific keywords
        key_matches = sum(1 for kw in domain_data[domain]["keywords"] if kw in text_lower)
        score = float(cos_scores[i]) + key_matches * keyword_boost

        if score > 0.6:
            scores.append((domain, 1))
        elif score >= 0.5:
            scores.append((domain, score))

    #Return top_k the most relevant industries
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]

#goes through all CVs, extracts scores of relevant domains based on the "project experience" section and saves a dictionary
#in the dictionary we have a specific domain and a certain score assigned
def compute_domain_scores_for_each_cv(cv_files, domain_data, cv_project_experiences):
    model = SentenceTransformer("all-MiniLM-L6-v2")  #fast & accurate
    cv_domain_experience_scores = {}

    for index in range(len(cv_files)):
        results = match_domains(model, domain_data, cv_project_experiences[index])
        if len(results) != 0:
            cv_domain_experience_scores[cv_files[index]] = {}
            for domain, score in results:
                cv_domain_experience_scores[cv_files[index]][domain] = score
    return cv_domain_experience_scores


#instead of calculating the scores each time, you load a saved .json (MEMOIZATION).
def get_domain_matching_scores(domain_selected, cv_files):
    with open("./industry/cv_domain_experience_scores.json", "r") as f:
        cv_domain_experience_scores = json.load(f)
        scores = [cv_domain_experience_scores.get(cv_file, {}).get(domain_selected, 0) for cv_file in cv_files]
        return np.array(scores) #return the score of each CV against the selected industry

# ------------------------------------------------------


#2---------------------------
def get_keyword_matching_scores(custom_skills, cv_text):
    #compute weighted match
    #receive a list of skills with weights
    skill_matches = {}
    for skill, weight in custom_skills:
        count = len(re.findall(rf"\b{re.escape(skill.lower())}\b", cv_text.lower()))
        skill_matches[skill] = (count > 0, weight) #stores in a dictionary whether it was found in the CV or not, and what weight it has
    #only add the weights of the skills found
    total_score = sum(weight for matched, weight in skill_matches.values() if matched)
    return total_score / 100

#for each CV, calculate the skill matching score and return a NumPy array with scores
def get_cv_keyword_matching_scores(custom_skills, cv_texts):
    scores = [get_keyword_matching_scores(custom_skills, cv_text) for cv_text in cv_texts]
    return np.array(scores)

#explanatory report for the score from the previous function
#improvised momentarily to which we will intervene with LLM for more complex explanations
def get_skills_matched_for_cv(custom_skills, cv_text):
    # Compute weighted match
    skill_matches = {}
    for skill, weight in custom_skills:
        count = len(re.findall(rf"\b{re.escape(skill.lower())}\b", cv_text.lower()))
        skill_matches[skill] = (count > 0, weight)
    return skill_matches
#---------------------------



#3---------------------------
# Custom tokenizer using spaCy
#takes a text, passes it through the SpaCy pipeline and returns a list of normalized words (lemmatized, lowercase),
#excluding stopwords, punctuation, spaces and numbers
def spacy_tokenizer(text):
    doc = nlp(text) #
    return [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and not token.text.isspace() and not token.text.isnumeric()]


def progress_bar_update(percent_complete, 
                        progress_bar,
                        status_text):
    progress_bar.progress(percent_complete)
    progress_text = "Operation in progress. Please wait. ⏳" if percent_complete < 100 else "Operation Completed."
    status_text.text(f"{progress_text} - {percent_complete}% - ")


def get_matching_scores_between_cvs_and_job_description(cv_texts, job_text, progress_bar, status_text):
    #all CVs and jobs are tokenized and cleaned with spacy_tokenizerPreprocesare
    #PREPROCESSING
    cv_texts_preprocessed = [" ".join(spacy_tokenizer(cv_text)) for cv_text in cv_texts]
    job_text_preprocessed = " ".join(spacy_tokenizer(job_text))

    progress_bar_update(30, progress_bar, status_text) #ui

    #SEMATIC SIMILARITY (SpaCy embeddings)
        #we use the embedding vectors from SpaCy (word Vectors / transformer Model). with the result being a value between 0 and 1.
            # => part that covers the semantic analysis part of the 60% requirement.
    similarities_embeddings = [nlp(job_text_preprocessed).similarity(nlp(cv_text_preprocessed)) for cv_text_preprocessed in cv_texts_preprocessed]

    progress_bar_update(70, progress_bar, status_text) #ui

    #TF-IDF similarity
    combined_cv_job_description_preprocessed = cv_texts_preprocessed + [job_text_preprocessed] #all texts (CV + job) are combined into a list

    # Fit vectorizer on combined
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(combined_cv_job_description_preprocessed)
    #sparse matrices with scores based on term frequency
    cv_vectors = tfidf_matrix[:len(cv_texts_preprocessed)]
    job_vectors = tfidf_matrix[len(cv_texts_preprocessed):]

    similarities_tfidf = cosine_similarity(job_vectors, cv_vectors) #cosine similarity between the job and each CV is calculated based on the TF-IDF vectors.
    #cos(0) = 1 => similarity ; cos(90) = 0 => no similarity
    progress_bar_update(80, progress_bar, status_text)

    #normalizes TF-IDF scores to [0, 1] so that they can be combined with SpaCy scores
    scaler = MinMaxScaler()    
    similarities_tfidf_scaled = scaler.fit_transform(np.array(similarities_tfidf).reshape(-1, 1)).flatten()
    
    progress_bar_update(90, progress_bar, status_text)
    similarities_embeddings = np.array(similarities_embeddings)

    final_embeddings_tfidf_scores = 0.6 * similarities_embeddings + 0.4 * similarities_tfidf_scaled

    #"""The system generates a match score (0–100%) based on the relevance of the developer's experience and skills to the job requirements."""
    return final_embeddings_tfidf_scores
#---------------------------




st.title("📌 Match Job → CVs")

st.header("Select domain matching candidates")
domain_selected = st.selectbox("What domain would you like the candidate to have experience in?", options=list(domain_data.keys()))

#let user define custom skills
st.subheader("🛠️ Define Required Skills and Weights")
skill_count = st.number_input("How many skills do you want to assess?", min_value=1, max_value=20, value=5)

custom_skills = []
total_weight = 0

if skill_count:
    with st.form("skill_form"):
        for i in range(skill_count):
            col1, col2 = st.columns([2, 1])
            with col1:
                skill_name = st.text_input(f"Skill #{i+1} name", key=f"skill_name_{i}")
            with col2:
                weight = st.number_input(f"Weight % for Skill #{i+1}", min_value=0, max_value=100, key=f"skill_weight_{i}")

            if skill_name:
                custom_skills.append((skill_name.strip(), weight))
                total_weight += weight
        st.markdown(f"**Total Weight: {total_weight}%**")
        submitted = st.form_submit_button("✅ Done selecting skills")

uploaded_file = st.file_uploader("Upload a Job Description (.docx format)", type=["docx"])
progress_bar = st.progress(0)
status_text = st.empty()

start_button = st.button("Get Matching CVs")

if start_button:
    if not uploaded_file:
        st.error("No job description uploaded (docx)")
    elif not domain_selected:
        st.error("No domain selected")
    elif total_weight != 100:
        st.error("Total weight must be exactly 100%.")
    else:
        #job_text = uploaded_file.read().decode("utf-8")
        doc = Document(uploaded_file)

        progress_bar_update(0, progress_bar, status_text)

        text = ' '.join([para.text for para in doc.paragraphs])
        text = text.replace('\n', ' ')
        text = text.replace('  ', ' ')
        job_text = text.split("Benefits:")[0]
        cv_texts, cv_files, cv_project_experiences = load_docx_from_folder(cv_folder)

        progress_bar_update(10, progress_bar, status_text)
        
        nlp = spacy.load("en_core_web_lg")
        #st.write(cv_texts)

        industry_knowledge_scores = get_domain_matching_scores(domain_selected, cv_files)
        
        technical_qualification_scores = get_cv_keyword_matching_scores(custom_skills, cv_texts)

        job_cv_matching_scores = get_matching_scores_between_cvs_and_job_description(cv_texts, job_text, progress_bar, status_text)

        final_scores = 0.1 * industry_knowledge_scores + 0.3 * technical_qualification_scores + 0.6 * job_cv_matching_scores

        final_scores_indices = np.argsort(final_scores)[::-1][:5] #sort descending and return indices for the first 5 CVs
        for i in final_scores_indices:
            st.write(f"{cv_files[i]} - Matching score: {final_scores[i] * 100:.2f}%")

        doc = Document(cv_folder + '/' + cv_files[final_scores_indices[0]])
        best_cv = '\n'.join([para.text for para in doc.paragraphs])

        st.header("The top matching CV is:")
        st.write(best_cv)

        st.header("Selection Explanation:")

        selected_cv_index = final_scores_indices[0]
        st.write(f"Industry knowledge score: {industry_knowledge_scores[selected_cv_index] * 100:.2f}% | Technical Qualification score: {technical_qualification_scores[selected_cv_index] * 100:.2f}% | Job - CV Matching score: {job_cv_matching_scores[selected_cv_index] * 100:.2f}%")
        skills_matched = get_skills_matched_for_cv(custom_skills, cv_texts[selected_cv_index])
        st.write(f"Technical Skills matched: {skills_matched}")
        #model = SentenceTransformer("all-MiniLM-L6-v2") 
        #match_domains(model, domain_data, cv_project_experiences[selected_cv_index])

        #get_explanations_for_top_selected_cv()
        progress_bar_update(100, progress_bar, status_text)
        st.balloons()