import streamlit as st
import spacy
import os
import numpy as np
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import MinMaxScaler
from sentence_transformers import SentenceTransformer, util
import re
import html
import time

# === CONFIG ===
cv_folder = './DataSet/cv'
job_folder = './DataSet/job_descriptions'

# === DOMAIN DATA ===
domain_data = { 
    "Banking": { "desc": "Developed software for banking platforms, digital wallets, loan processing, or credit systems.", 
                 "keywords": ["bank", "loan", "credit", "atm", "fintech", "interest", "account", "ledger"] }, 
    "Healthcare": { "desc": "Built medical systems such as EHR, hospital platforms, clinical apps, or telemedicine services.", 
                    "keywords": ["healthcare", "ehr", "patient", "hospital", "clinic", "medical", "doctor", "nurse"] }, 
    "E-commerce": { "desc": "Created platforms or tools for online stores, shopping carts, payments, or product discovery.", 
                    "keywords": ["ecommerce", "checkout", "cart", "payment", "shopify", "woocommerce", "product", "sku"] }, 
    "Telecommunications": { "desc": "Engineered tools for telecom networks, call management, VoIP, or network monitoring.", 
                            "keywords": ["telecom", "sms", "voip", "5g", "network", "bandwidth", "subscriber", "lte"] }, 
    "Education": { "desc": "Built education platforms such as learning portals, student dashboards, or LMS systems.", 
                   "keywords": ["education", "student", "teacher", "learning", "course", "classroom", "lms", "school"] }, 
    "Retail": { "desc": "Developed software for retail businesses such as POS systems, inventory, or loyalty programs.", 
                "keywords": ["retail", "store", "inventory", "pos", "stock", "sku", "receipt", "shopping"] }, 
    "Insurance": { "desc": "Created applications for policy management, claims processing, or underwriting systems.", 
                   "keywords": ["insurance", "claim", "policy", "underwriting", "premium", "broker", "risk"] }, 
    "Legal": { "desc": "Built tools for legal case management, document processing, or e-discovery platforms.", 
               "keywords": ["legal", "law", "contract", "case", "compliance", "jurisdiction", "litigation"] }, 
    "Manufacturing": { "desc": "Developed automation systems, supply chain tools, or MES solutions for production plants.", 
                       "keywords": ["manufacturing", "plant", "automation", "mes", "machine", "factory", "assembly"] }, 
    "Transportation & Logistics": { "desc": "Built platforms for delivery tracking, fleet management, logistics optimization, or routing.", 
                                    "keywords": ["logistics", "delivery", "routing", "fleet", "dispatch", "transport", "warehouse"] }, 
    "Energy & Utilities": { "desc": "Developed monitoring systems, SCADA platforms, or analytics for power and water utilities.", 
                            "keywords": ["energy", "power", "electricity", "gas", "grid", "meter", "solar", "utility", "scada"] }, 
    "Real Estate": { "desc": "Engineered platforms for property listings, CRM tools for agents, or real estate analytics.", 
                     "keywords": ["real estate", "property", "mortgage", "agent", "listing", "tenant", "lease"] }, 
    "Government": { "desc": "Built public service portals, civic data dashboards, or digital identity platforms.", 
                    "keywords": ["government", "municipal", "civic", "permit", "id", "citizen", "registry"] }, 
    "Marketing": { "desc": "Built marketing analytics tools, email campaign platforms, or digital ad performance systems.", 
                   "keywords": ["marketing", "campaign", "seo", "email", "ads", "promotion", "branding", "targeting"] }, 
    "Media & Entertainment": { "desc": "Created digital content platforms, streaming services, or entertainment production tools.", 
                               "keywords": ["media", "streaming", "video", "music", "entertainment", "broadcast", "subscriber"] }, 
    "Construction": { "desc": "Engineered project management tools, BIM integrations, or field apps for construction teams.", 
                      "keywords": ["construction", "site", "blueprint", "project", "bim", "architect", "contractor"] }, 
    "Finance (non-banking)": { "desc": "Worked on accounting systems, budgeting tools, payroll, or financial planning platforms.", 
                               "keywords": ["finance", "accounting", "budget", "payroll", "invoice", "expense", "audit", "report"] }
}

# === UTILITY FUNCTIONS ===
def extract_text_from_docx(filepath, is_cv):
    doc = Document(filepath)
    return ' '.join([para.text for para in doc.paragraphs[1 if is_cv else 0:]])

def load_docx_from_folder(folder_path, is_cv=True):
    documents, filenames, descriptions = [], [], []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            full_path = os.path.join(folder_path, filename)
            try:
                text = extract_text_from_docx(full_path, is_cv)
                text = text.replace('\n', ' ').replace('  ', ' ').strip()
                if not text:
                    print(f"⚠️ File '{filename}' is empty after cleaning. Skipping.")
                    continue
                if is_cv:
                    description = text.split("Project Experience")[1] if "Project Experience" in text else text
                else:
                    description = text.split("Key Responsibilities:")[0] if "Key Responsibilities:" in text else text
                documents.append(text)
                filenames.append(filename)
                descriptions.append(description)
            except Exception as e:
                print(f"⚠️ Could not read '{filename}': {e}")
    return documents, filenames, descriptions

def spacy_tokenizer(text):
    doc = nlp(text)
    return [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and not token.text.isspace() and not token.text.isnumeric()]

def match_domains(model, domain_data, text, top_k=5, keyword_boost=0.1):
    domain_names = list(domain_data.keys())
    domain_embeddings = model.encode([domain_data[d]["desc"] for d in domain_names], convert_to_tensor=True)
    text_embedding = model.encode(text, convert_to_tensor=True)
    cos_scores = util.cos_sim(text_embedding, domain_embeddings)[0]
    scores = []
    text_lower = text.lower()
    for i, domain in enumerate(domain_names):
        key_matches = sum(1 for kw in domain_data[domain]["keywords"] if kw in text_lower)
        score = float(cos_scores[i]) + key_matches * keyword_boost
        if score > 0.6:
            scores.append((domain, 1))
        elif score >= 0.5:
            scores.append((domain, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]

def get_matching_scores_between_job_descriptions_and_cv(job_texts, cv_text):
    cv_preprocessed = " ".join(spacy_tokenizer(cv_text))
    job_texts_preprocessed = [" ".join(spacy_tokenizer(job)) for job in job_texts]
    similarities_embeddings = [nlp(cv_preprocessed).similarity(nlp(job)) for job in job_texts_preprocessed]
    combined = [cv_preprocessed] + job_texts_preprocessed
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(combined)
    cv_vector = tfidf_matrix[0]
    job_vectors = tfidf_matrix[1:]
    similarities_tfidf = cosine_similarity(cv_vector, job_vectors)[0]
    scaler = MinMaxScaler()
    similarities_tfidf_scaled = scaler.fit_transform(similarities_tfidf.reshape(-1, 1)).flatten()
    final_scores = 0.6 * np.array(similarities_embeddings) + 0.4 * similarities_tfidf_scaled
    return final_scores

def progress_bar_update(percent_complete, progress_bar, status_text):
    progress_bar.progress(percent_complete)
    progress_text = "Processing... ⏳" if percent_complete < 100 else "Done. ✅"
    status_text.text(f"{progress_text} {percent_complete}%")

# === UI ===
st.set_page_config(page_title="Match CV to Job", layout="centered")
st.markdown("""
    <style>
    html, body, [class*="css"]  {
        font-family: 'Segoe UI', sans-serif;
        color: #222;
        background-color: #f4f6f9;
    }
    </style>
""", unsafe_allow_html=True)
st.title("🎯 CV to Job Matching System")

st.markdown("""
Upload your **CV in .docx format**, and we'll match it against available job descriptions.
The system uses semantic similarity and industry domain matching to determine the best fit.
""")

nlp = spacy.load("en_core_web_lg")
model = SentenceTransformer("all-MiniLM-L6-v2")

uploaded_file = st.file_uploader("📤 Upload your CV (.docx only)", type=["docx"])
progress_bar = st.progress(0)
status_text = st.empty()

if st.button("🔍 Find Best Match"):
    if not uploaded_file:
        st.error("⚠️ Please upload a CV file before proceeding.")
    else:
        text = ' '.join([para.text for para in Document(uploaded_file).paragraphs])
        text = text.replace('\n', ' ').replace('  ', ' ')
        cv_text = text
        cv_description = text.split("Project Experience")[1] if "Project Experience" in text else text

        progress_bar_update(10, progress_bar, status_text)
        job_texts, job_files, job_descriptions = load_docx_from_folder(job_folder, is_cv=False)

        if not job_texts:
            st.error("❌ No valid job descriptions found. Please upload at least one job description file.")
            st.stop()

        domain_scores = match_domains(model, domain_data, cv_description)
        top_domain = domain_scores[0][0] if domain_scores else "Unknown"

        industry_scores = np.array([domain_scores[0][1] if top_domain.lower() in f.lower() else 0 for f in job_files])
        st.info(f"🔍 **Detected domain expertise**: `{top_domain}`")
        progress_bar_update(40, progress_bar, status_text)

        job_match_scores = get_matching_scores_between_job_descriptions_and_cv(job_texts, cv_text)
        progress_bar_update(80, progress_bar, status_text)

        final_scores = 0.1 * industry_scores + 0.9 * job_match_scores
        top_match_index = np.argmax(final_scores)
        progress_bar_update(100, progress_bar, status_text)

        st.success("✅ Best Matching Job Identified")
        placeholder = st.empty()
        placeholder.markdown("⏳ Matching complete. Rendering results...")
        time.sleep(1.5)
        placeholder.empty()

        st.metric(label="🎯 Job Title", value=job_files[top_match_index])
        st.metric(label="📈 Match Score", value=f"{final_scores[top_match_index]*100:.2f}%")

        st.markdown("### 📊 Matching Breakdown")

        st.markdown(f"""
            <div style="margin-bottom:1rem;">
                <span style="display:inline-block; background:#e8f5e9; color:#2e7d32; padding:0.4rem 0.8rem; border-radius:8px; font-weight:bold;">
                🏷️ Domain Match: {industry_scores[top_match_index]*100:.2f}%
                </span>
            </div>
            <div>
                <span style="display:inline-block; background:#e3f2fd; color:#1565c0; padding:0.4rem 0.8rem; border-radius:8px; font-weight:bold;">
                🧠 Semantic Similarity: {job_match_scores[top_match_index]*100:.2f}%
                </span>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("### 📝 Job Description Preview")
        best_doc = Document(os.path.join(job_folder, job_files[top_match_index]))
        job_text = '\n'.join([para.text for para in best_doc.paragraphs])
        with st.expander("📄 View Full Job Description"):
            st.write(job_text)
        st.download_button(
            label="💾 Download Job Description (.txt)",
            data=job_text,
            file_name=job_files[top_match_index].replace(".docx", ".txt"),
            mime="text/plain"
            )

        st.balloons()
